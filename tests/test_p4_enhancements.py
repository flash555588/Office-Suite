"""P4 DOCX/XLSX 功能增强测试

验证新增功能：
1. DOCX：页边距、列表、段落间距、表格表头样式、data_ref
2. XLSX：数字格式、图表坐标轴、柱状图方向
"""

import sys
from pathlib import Path

PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from office_suite.dsl.parser import parse_yaml_string
from office_suite.ir.compiler import compile_document
from office_suite.renderer.docx.document import DOCXRenderer
from office_suite.renderer.xlsx.workbook import XLSXRenderer

OUTPUT_DIR = PROJECT_ROOT / "tests" / "output"
OUTPUT_DIR.mkdir(exist_ok=True)


# ============================================================
# DOCX 功能增强测试
# ============================================================

def test_docx_page_margins():
    """DOCX 页边距设置为 25mm"""
    from docx import Document as DocxDocument

    dsl = """
version: "4.0"
type: document
slides:
  - layout: blank
    elements:
      - type: text
        content: "测试"
        position: { x: 20mm, y: 20mm, width: 100mm, height: 10mm }
"""
    doc = parse_yaml_string(dsl)
    ir = compile_document(doc)
    output = OUTPUT_DIR / "p4_margins.docx"
    DOCXRenderer().render(ir, output)

    docx = DocxDocument(str(output))
    section = docx.sections[0]
    # 25mm ≈ 900000 EMU (python-docx 可能有微小误差)
    assert abs(section.top_margin - 900000) < 5000, f"top_margin={section.top_margin}"
    assert abs(section.left_margin - 900000) < 5000, f"left_margin={section.left_margin}"


def test_docx_bullet_list():
    """DOCX 无序列表"""
    from docx import Document as DocxDocument

    dsl = """
version: "4.0"
type: document
slides:
  - layout: blank
    elements:
      - type: text
        content: "第一项"
        extra: { list_type: bullet }
        position: { x: 20mm, y: 20mm, width: 200mm, height: 8mm }
      - type: text
        content: "第二项"
        extra: { list_type: bullet }
        position: { x: 20mm, y: 30mm, width: 200mm, height: 8mm }
"""
    doc = parse_yaml_string(dsl)
    ir = compile_document(doc)
    output = OUTPUT_DIR / "p4_bullet.docx"
    DOCXRenderer().render(ir, output)

    docx = DocxDocument(str(output))
    list_paras = [p for p in docx.paragraphs if "List" in p.style.name]
    assert len(list_paras) >= 2, f"列表段落数 {len(list_paras)}"
    assert "第一项" in list_paras[0].text
    assert "第二项" in list_paras[1].text


def test_docx_numbered_list():
    """DOCX 有序列表"""
    from docx import Document as DocxDocument

    dsl = """
version: "4.0"
type: document
slides:
  - layout: blank
    elements:
      - type: text
        content: "步骤一"
        extra: { list_type: number }
        position: { x: 20mm, y: 20mm, width: 200mm, height: 8mm }
      - type: text
        content: "步骤二"
        extra: { list_type: number }
        position: { x: 20mm, y: 30mm, width: 200mm, height: 8mm }
"""
    doc = parse_yaml_string(dsl)
    ir = compile_document(doc)
    output = OUTPUT_DIR / "p4_numbered.docx"
    DOCXRenderer().render(ir, output)

    docx = DocxDocument(str(output))
    list_paras = [p for p in docx.paragraphs if "Number" in p.style.name]
    assert len(list_paras) >= 2, f"有序列表段落数 {len(list_paras)}"


def test_docx_table_header_style():
    """DOCX 表头深色背景 + 居中 + 白色文字"""
    from docx import Document as DocxDocument

    dsl = """
version: "4.0"
type: document
slides:
  - layout: blank
    elements:
      - type: table
        rows: 2
        cols: 2
        data:
          - ["名称", "值"]
          - ["测试", 42]
        position: { x: 20mm, y: 20mm, width: 100mm, height: 30mm }
"""
    doc = parse_yaml_string(dsl)
    ir = compile_document(doc)
    output = OUTPUT_DIR / "p4_table_header.docx"
    DOCXRenderer().render(ir, output)

    docx = DocxDocument(str(output))
    table = docx.tables[0]

    # 表头单元格应有深色背景
    cell = table.cell(0, 0)
    tc_pr = cell._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
    assert tc_pr is not None, "表头无背景色"


def test_docx_data_ref():
    """DOCX 表格 data_ref 引用"""
    from docx import Document as DocxDocument

    dsl = """
version: "4.0"
type: document
data:
  my_table:
    inline:
      - ["项目", "状态"]
      - ["A", "完成"]
      - ["B", "进行中"]
slides:
  - layout: blank
    elements:
      - type: table
        data_ref: my_table
        rows: 3
        cols: 2
        position: { x: 20mm, y: 20mm, width: 100mm, height: 40mm }
"""
    doc = parse_yaml_string(dsl)
    ir = compile_document(doc)
    output = OUTPUT_DIR / "p4_data_ref.docx"
    DOCXRenderer().render(ir, output)

    docx = DocxDocument(str(output))
    table = docx.tables[0]
    assert table.cell(1, 0).text == "A"
    assert table.cell(2, 1).text == "进行中"


# ============================================================
# XLSX 功能增强测试
# ============================================================

def test_xlsx_number_format():
    """XLSX 数字格式"""
    from openpyxl import load_workbook

    dsl = """
version: "4.0"
type: spreadsheet
slides:
  - layout: blank
    elements:
      - type: table
        rows: 3
        cols: 2
        data:
          - ["项目", "金额"]
          - ["A", 1234.5]
          - ["B", 5678.9]
        extra: { number_format: currency }
        position: { x: 0mm, y: 0mm }
"""
    doc = parse_yaml_string(dsl)
    ir = compile_document(doc)
    output = OUTPUT_DIR / "p4_number_fmt.xlsx"
    XLSXRenderer().render(ir, output)

    wb = load_workbook(str(output))
    ws = wb.active
    # 数据行应有货币格式
    cell = ws.cell(2, 2)
    assert "¥" in cell.number_format or cell.number_format == '¥#,##0.00', f"格式: {cell.number_format}"


def test_xlsx_chart_axis_labels():
    """XLSX 图表坐标轴标签"""
    from openpyxl import load_workbook

    dsl = """
version: "4.0"
type: spreadsheet
slides:
  - layout: blank
    elements:
      - type: chart
        chart_type: bar
        extra:
          title: "测试"
          categories: ["A", "B"]
          series:
            - name: "系列1"
              values: [10, 20]
          x_axis_title: "类别"
          y_axis_title: "数值"
        position: { x: 0mm, y: 0mm }
"""
    doc = parse_yaml_string(dsl)
    ir = compile_document(doc)
    output = OUTPUT_DIR / "p4_chart_axis.xlsx"
    XLSXRenderer().render(ir, output)

    wb = load_workbook(str(output))
    ws = wb.active
    chart = ws._charts[0]
    # openpyxl axis.title 是 Title 对象，需通过 tx 获取文本
    x_text = chart.x_axis.title.tx.rich.paragraphs[0].r[0].t
    y_text = chart.y_axis.title.tx.rich.paragraphs[0].r[0].t
    assert x_text == "类别"
    assert y_text == "数值"


def test_xlsx_column_chart_type():
    """XLSX 柱状图方向（column = 垂直）"""
    from openpyxl import load_workbook

    dsl = """
version: "4.0"
type: spreadsheet
slides:
  - layout: blank
    elements:
      - type: chart
        chart_type: column
        extra:
          categories: ["A", "B"]
          series:
            - name: "S1"
              values: [1, 2]
        position: { x: 0mm, y: 0mm }
"""
    doc = parse_yaml_string(dsl)
    ir = compile_document(doc)
    output = OUTPUT_DIR / "p4_col_chart.xlsx"
    XLSXRenderer().render(ir, output)

    wb = load_workbook(str(output))
    ws = wb.active
    chart = ws._charts[0]
    assert chart.type == "col", f"图表类型: {chart.type}"


def test_xlsx_percent_format():
    """XLSX 百分比格式"""
    from openpyxl import load_workbook

    dsl = """
version: "4.0"
type: spreadsheet
slides:
  - layout: blank
    elements:
      - type: table
        rows: 2
        cols: 2
        data:
          - ["项目", "完成率"]
          - ["A", 0.85]
        extra: { number_format: percent }
        position: { x: 0mm, y: 0mm }
"""
    doc = parse_yaml_string(dsl)
    ir = compile_document(doc)
    output = OUTPUT_DIR / "p4_percent.xlsx"
    XLSXRenderer().render(ir, output)

    wb = load_workbook(str(output))
    ws = wb.active
    cell = ws.cell(2, 2)
    assert "%" in cell.number_format, f"格式: {cell.number_format}"


if __name__ == "__main__":
    import pytest
    pytest.main([__file__, "-v"])
