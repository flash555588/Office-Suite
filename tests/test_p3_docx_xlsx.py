"""P3 DOCX/XLSX 内部结构验证测试

验证生成的文档内部结构正确性，而不只是文件大小。

覆盖：
1. DOCX：段落内容、标题层级、表格维度、单元格值、图片
2. XLSX：Sheet 名称、单元格值、图表对象、列宽
"""

import sys
from pathlib import Path

PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from office_suite.dsl.parser import parse_yaml_string
from office_suite.ir.compiler import compile_document
from office_suite.renderer.docx.document import DOCXRenderer
from office_suite.renderer.xlsx.workbook import XLSXRenderer

OUTPUT_DIR = PROJECT_ROOT / "tests" / "output"
OUTPUT_DIR.mkdir(exist_ok=True)


# ============================================================
# DOCX 内部结构测试
# ============================================================

def _render_docx(dsl: str, name: str):
    """解析 DSL 并渲染为 DOCX，返回文件路径"""
    doc = parse_yaml_string(dsl)
    ir = compile_document(doc)
    output = OUTPUT_DIR / f"p3_{name}.docx"
    renderer = DOCXRenderer()
    return renderer.render(ir, output)


def test_docx_paragraph_content():
    """DOCX 段落内容正确"""
    from docx import Document as DocxDocument

    dsl = """
version: "4.0"
type: document
slides:
  - layout: blank
    elements:
      - type: text
        content: "测试标题"
        style: { font: { size: 32, weight: 700 } }
        position: { x: 20mm, y: 20mm, width: 200mm, height: 15mm }
      - type: text
        content: "这是正文内容"
        style: { font: { size: 14 } }
        position: { x: 20mm, y: 40mm, width: 200mm, height: 10mm }
"""
    path = _render_docx(dsl, "para_content")
    doc = DocxDocument(str(path))

    # 至少有 2 个段落（标题 + 正文）
    assert len(doc.paragraphs) >= 2, f"段落数 {len(doc.paragraphs)} < 2"

    # 检查标题
    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert len(headings) >= 1, "没有标题段落"
    assert "测试标题" in headings[0].text

    # 检查正文
    body = [p for p in doc.paragraphs if not p.style.name.startswith("Heading")]
    body_texts = [p.text for p in body]
    assert any("这是正文内容" in t for t in body_texts), f"正文内容未找到: {body_texts}"


def test_docx_heading_level():
    """DOCX 标题层级正确（font_size >= 28 → H1, >= 20 → H2）"""
    from docx import Document as DocxDocument

    dsl = """
version: "4.0"
type: document
slides:
  - layout: blank
    elements:
      - type: text
        content: "一级标题"
        style: { font: { size: 32 } }
        position: { x: 20mm, y: 20mm, width: 200mm, height: 15mm }
      - type: text
        content: "二级标题"
        style: { font: { size: 22 } }
        position: { x: 20mm, y: 40mm, width: 200mm, height: 12mm }
      - type: text
        content: "正文"
        style: { font: { size: 14 } }
        position: { x: 20mm, y: 60mm, width: 200mm, height: 10mm }
"""
    path = _render_docx(dsl, "heading_level")
    doc = DocxDocument(str(path))

    h1 = [p for p in doc.paragraphs if p.style.name == "Heading 1"]
    h2 = [p for p in doc.paragraphs if p.style.name == "Heading 2"]

    assert len(h1) >= 1, "没有 Heading 1"
    assert "一级标题" in h1[0].text
    assert len(h2) >= 1, "没有 Heading 2"
    assert "二级标题" in h2[0].text


def test_docx_table_dimensions():
    """DOCX 表格维度和内容正确"""
    from docx import Document as DocxDocument

    dsl = """
version: "4.0"
type: document
slides:
  - layout: blank
    elements:
      - type: table
        rows: 3
        cols: 2
        data:
          - ["姓名", "年龄"]
          - ["张三", "25"]
          - ["李四", "30"]
        position: { x: 20mm, y: 20mm, width: 100mm, height: 50mm }
"""
    path = _render_docx(dsl, "table_dims")
    doc = DocxDocument(str(path))

    assert len(doc.tables) >= 1, "没有表格"
    table = doc.tables[0]

    assert len(table.rows) == 3, f"行数 {len(table.rows)} != 3"
    assert len(table.columns) == 2, f"列数 {len(table.columns)} != 2"

    # 检查单元格值
    assert table.cell(0, 0).text == "姓名"
    assert table.cell(0, 1).text == "年龄"
    assert table.cell(1, 0).text == "张三"
    assert table.cell(2, 1).text == "30"


def test_docx_table_header_bold():
    """DOCX 表格首行加粗"""
    from docx import Document as DocxDocument

    dsl = """
version: "4.0"
type: document
slides:
  - layout: blank
    elements:
      - type: table
        rows: 2
        cols: 2
        data:
          - ["列A", "列B"]
          - ["值1", "值2"]
        position: { x: 20mm, y: 20mm, width: 100mm, height: 30mm }
"""
    path = _render_docx(dsl, "table_bold")
    doc = DocxDocument(str(path))

    table = doc.tables[0]
    # 首行单元格的 run 应该加粗
    cell = table.cell(0, 0)
    for para in cell.paragraphs:
        for run in para.runs:
            assert run.bold is True, f"表头未加粗: {run.text}"


def test_docx_image_placeholder():
    """DOCX 图片缺失时降级为占位符"""
    from docx import Document as DocxDocument

    dsl = """
version: "4.0"
type: document
slides:
  - layout: blank
    elements:
      - type: image
        source: "nonexistent.png"
        position: { x: 20mm, y: 20mm, width: 100mm, height: 80mm }
"""
    path = _render_docx(dsl, "img_placeholder")
    doc = DocxDocument(str(path))

    # 应该有占位符段落
    texts = [p.text for p in doc.paragraphs]
    assert any("图片" in t or "nonexistent" in t for t in texts), f"图片占位符未找到: {texts}"


# ============================================================
# XLSX 内部结构测试
# ============================================================

def _render_xlsx(dsl: str, name: str):
    """解析 DSL 并渲染为 XLSX，返回文件路径"""
    doc = parse_yaml_string(dsl)
    ir = compile_document(doc)
    output = OUTPUT_DIR / f"p3_{name}.xlsx"
    renderer = XLSXRenderer()
    return renderer.render(ir, output)


def test_xlsx_sheet_names():
    """XLSX Sheet 名称 — 多 Sheet 自动生成"""
    from openpyxl import load_workbook

    dsl = """
version: "4.0"
type: spreadsheet
slides:
  - layout: blank
    elements:
      - type: table
        rows: 2
        cols: 2
        data: [["A","B"],["1","2"]]
        position: { x: 0mm, y: 0mm }
  - layout: blank
    elements:
      - type: table
        rows: 2
        cols: 2
        data: [["C","D"],["3","4"]]
        position: { x: 0mm, y: 0mm }
"""
    path = _render_xlsx(dsl, "sheet_names")
    wb = load_workbook(str(path))

    # 当前行为：自动生成 Sheet1, Sheet2
    assert len(wb.sheetnames) == 2, f"Sheet 数: {len(wb.sheetnames)}"
    assert "Sheet1" in wb.sheetnames
    assert "Sheet2" in wb.sheetnames


def test_xlsx_cell_values():
    """XLSX 单元格值正确"""
    from openpyxl import load_workbook

    dsl = """
version: "4.0"
type: spreadsheet
slides:
  - layout: blank
    elements:
      - type: table
        rows: 3
        cols: 2
        data:
          - ["产品", "价格"]
          - ["苹果", 5]
          - ["香蕉", 3]
        position: { x: 0mm, y: 0mm }
"""
    path = _render_xlsx(dsl, "cell_values")
    wb = load_workbook(str(path))
    ws = wb.active

    assert ws.cell(1, 1).value == "产品"
    assert ws.cell(1, 2).value == "价格"
    assert ws.cell(2, 1).value == "苹果"
    assert ws.cell(2, 2).value == 5
    assert ws.cell(3, 1).value == "香蕉"
    assert ws.cell(3, 2).value == 3


def test_xlsx_header_style():
    """XLSX 表头样式（加粗、深色背景）"""
    from openpyxl import load_workbook

    dsl = """
version: "4.0"
type: spreadsheet
slides:
  - layout: blank
    elements:
      - type: table
        rows: 2
        cols: 2
        data:
          - ["名称", "值"]
          - ["测试", 42]
        position: { x: 0mm, y: 0mm }
"""
    path = _render_xlsx(dsl, "header_style")
    wb = load_workbook(str(path))
    ws = wb.active

    header_cell = ws.cell(1, 1)
    assert header_cell.font.bold is True, "表头未加粗"
    assert header_cell.fill.start_color.rgb is not None, "表头无背景色"


def test_xlsx_chart_exists():
    """XLSX 图表对象存在"""
    from openpyxl import load_workbook

    dsl = """
version: "4.0"
type: spreadsheet
slides:
  - layout: blank
    elements:
      - type: chart
        chart_type: bar
        extra:
          title: "测试图表"
          categories: ["A", "B", "C"]
          series:
            - name: "系列1"
              values: [10, 20, 30]
        position: { x: 0mm, y: 0mm }
"""
    path = _render_xlsx(dsl, "chart_exists")
    wb = load_workbook(str(path))
    ws = wb.active

    assert len(ws._charts) >= 1, f"图表数 {len(ws._charts)} == 0"


def test_xlsx_chart_title():
    """XLSX 图表标题正确"""
    from openpyxl import load_workbook

    dsl = """
version: "4.0"
type: spreadsheet
slides:
  - layout: blank
    elements:
      - type: chart
        chart_type: line
        extra:
          title: "趋势图"
          categories: ["Q1", "Q2"]
          series:
            - name: "营收"
              values: [100, 200]
        position: { x: 0mm, y: 0mm }
"""
    path = _render_xlsx(dsl, "chart_title")
    wb = load_workbook(str(path))
    ws = wb.active

    chart = ws._charts[0]
    # openpyxl chart.title 是 Title 对象，需通过 tx 属性获取文本
    title_text = chart.title.tx.rich.paragraphs[0].r[0].t
    assert title_text == "趋势图", f"图表标题: {title_text}"


def test_xlsx_column_auto_width():
    """XLSX 自动列宽生效"""
    from openpyxl import load_workbook

    dsl = """
version: "4.0"
type: spreadsheet
slides:
  - layout: blank
    elements:
      - type: table
        rows: 2
        cols: 2
        data:
          - ["短", "这是一个很长的列标题"]
          - ["A", "B"]
        position: { x: 0mm, y: 0mm }
"""
    path = _render_xlsx(dsl, "col_width")
    wb = load_workbook(str(path))
    ws = wb.active

    # 第二列应该比第一列宽
    col_a = ws.column_dimensions["A"].width
    col_b = ws.column_dimensions["B"].width
    assert col_b > col_a, f"列宽 A={col_a}, B={col_b}"


def test_xlsx_multi_sheet_content():
    """XLSX 多 Sheet 各自内容独立"""
    from openpyxl import load_workbook

    dsl = """
version: "4.0"
type: spreadsheet
slides:
  - layout: blank
    extra: { title: "Sheet1" }
    elements:
      - type: table
        rows: 2
        cols: 1
        data: [["数据A"],["值1"]]
        position: { x: 0mm, y: 0mm }
  - layout: blank
    extra: { title: "Sheet2" }
    elements:
      - type: table
        rows: 2
        cols: 1
        data: [["数据B"],["值2"]]
        position: { x: 0mm, y: 0mm }
"""
    path = _render_xlsx(dsl, "multi_content")
    wb = load_workbook(str(path))

    ws1 = wb["Sheet1"]
    ws2 = wb["Sheet2"]

    assert ws1.cell(1, 1).value == "数据A"
    assert ws2.cell(1, 1).value == "数据B"
    assert ws1.cell(2, 1).value == "值1"
    assert ws2.cell(2, 1).value == "值2"


if __name__ == "__main__":
    import pytest
    pytest.main([__file__, "-v"])
